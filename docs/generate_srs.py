import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Styling and XML Helper Functions
# ---------------------------------------------------------------------------

def set_cell_background(cell, hex_color):
    """Sets the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_left_border(cell, hex_color, size="36"):
    """Sets a thick left border on a cell and removes top/bottom/right borders."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size) # 36 = 4.5pt
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), hex_color)
    tcBorders.append(left)
    
    for border_name in ['top', 'right', 'bottom']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
        
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding (margins) for a cell in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_run_font(run, name="Calibri", size_pt=11, bold=False, italic=False, color_rgb=None):
    """Sets font attributes on a text run."""
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def format_paragraph(paragraph, space_before=0, space_after=6, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Sets spacing and alignment for a paragraph."""
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.alignment = alignment

def add_styled_heading(doc, text, level):
    """Adds a custom-styled heading that matches premium design standards."""
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    
    if level == 1:
        format_paragraph(p, space_before=18, space_after=8)
        run = p.add_run(text)
        set_run_font(run, name="Calibri Light", size_pt=20, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x79))
    elif level == 2:
        format_paragraph(p, space_before=14, space_after=6)
        run = p.add_run(text)
        set_run_font(run, name="Calibri Light", size_pt=14, bold=True, color_rgb=RGBColor(0x2F, 0x55, 0x97))
    else:
        format_paragraph(p, space_before=10, space_after=4)
        run = p.add_run(text)
        set_run_font(run, name="Calibri", size_pt=12, bold=True, color_rgb=RGBColor(0x59, 0x59, 0x59))
    return p

def create_callout_box(doc, text, hex_bg="F2F4F7", hex_border="1F4E79"):
    """Creates a beautifully padded single-cell table acting as a callout box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, hex_bg)
    set_cell_left_border(cell, hex_border, size="36")
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    p = cell.paragraphs[0]
    format_paragraph(p, space_before=4, space_after=4, line_spacing=1.15)
    run = p.add_run(text)
    set_run_font(run, size_pt=10.5, italic=True, color_rgb=RGBColor(0x33, 0x33, 0x33))
    
    p_after = doc.add_paragraph()
    format_paragraph(p_after, space_before=0, space_after=6)

def create_code_block(doc, code_text):
    """Creates a grey-shaded code block table for raw PlantUML definitions."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F9FBFD")
    set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b_name in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{b_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4') # 0.5pt
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'D0D5DD')
        tcBorders.append(b)
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    format_paragraph(p, space_before=2, space_after=2, line_spacing=1.0)
    run = p.add_run(code_text)
    set_run_font(run, name="Courier New", size_pt=8.5, color_rgb=RGBColor(0x2B, 0x2B, 0x2B))
    
    p_after = doc.add_paragraph()
    format_paragraph(p_after, space_before=0, space_after=6)

def create_styled_table(doc, headers, data, col_widths=None):
    """Creates a beautifully styled table with custom widths, padded cells, and alternating colors."""
    tbl = doc.add_table(rows=len(data) + 1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    # Header Row
    hdr_cells = tbl.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "1F4E79")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        
        p = hdr_cells[i].paragraphs[0]
        format_paragraph(p, space_before=2, space_after=2, line_spacing=1.0)
        run = p.runs[0]
        set_run_font(run, name="Calibri", size_pt=10, bold=True, color_rgb=RGBColor(255, 255, 255))
        
        if col_widths and i < len(col_widths):
            hdr_cells[i].width = Inches(col_widths[i])
            
    # Set header row repeating
    trPr = tbl.rows[0]._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    trPr.append(tblHeader)
    
    # Data Rows
    for r_idx, row_data in enumerate(data):
        row_cells = tbl.rows[r_idx + 1].cells
        bg_color = "F7F9FB" if r_idx % 2 == 1 else "FFFFFF"
        
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=120, right=120)
            
            tcPr = row_cells[c_idx]._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            
            bottom_b = OxmlElement('w:bottom')
            bottom_b.set(qn('w:val'), 'single')
            bottom_b.set(qn('w:sz'), '4') # 0.5pt
            bottom_b.set(qn('w:color'), 'E5E5E5')
            tcBorders.append(bottom_b)
            
            for side in ['top', 'left', 'right']:
                side_b = OxmlElement(f'w:{side}')
                side_b.set(qn('w:val'), 'none')
                tcBorders.append(side_b)
            tcPr.append(tcBorders)
            
            p = row_cells[c_idx].paragraphs[0]
            format_paragraph(p, space_before=2, space_after=2, line_spacing=1.1)
            if p.runs:
                set_run_font(p.runs[0], name="Calibri", size_pt=9.5)
            
            if col_widths and c_idx < len(col_widths):
                row_cells[c_idx].width = Inches(col_widths[c_idx])
                
    p_after = doc.add_paragraph()
    format_paragraph(p_after, space_before=0, space_after=8)
    return tbl

def add_bullet_point(doc, text, bold_prefix=None):
    """Adds a clean, custom-spaced bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    format_paragraph(p, space_before=2, space_after=4, line_spacing=1.15)
    
    if bold_prefix:
        run_pref = p.add_run(bold_prefix)
        set_run_font(run_pref, bold=True)
        
    run_text = p.add_run(text)
    set_run_font(run_text)

def add_page_number_to_run(run):
    """Inserts a PAGE field XML tag inside a text run."""
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    run._r.append(fldSimple)

def setup_header_footer(doc):
    """Configures elegant headers and footers for all sections except the cover page."""
    for idx, section in enumerate(doc.sections):
        section.different_first_page_header_footer = True
        
        # Setup standard Header
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "Software Requirements Specification (SRS)  |  Money Manager Java Suite"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in hp.runs:
            set_run_font(run, name="Calibri", size_pt=8.5, color_rgb=RGBColor(128, 128, 128))
            
        # Setup standard Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = "Confidential  -  For Internal Development Use Only                                           Page "
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in fp.runs:
            set_run_font(run, name="Calibri", size_pt=8.5, color_rgb=RGBColor(128, 128, 128))
            
        run_num = fp.add_run()
        add_page_number_to_run(run_num)
        set_run_font(run_num, name="Calibri", size_pt=8.5, color_rgb=RGBColor(128, 128, 128))

def read_puml_file(file_path):
    """Reads PlantUML content from the local directory if it exists."""
    if os.path.exists(file_path):
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except Exception as e:
            return f"' Error reading file: {str(e)}"
    return f"' PlantUML source file not found: {os.path.basename(file_path)}"

# ---------------------------------------------------------------------------
# Main Document Generator
# ---------------------------------------------------------------------------

def generate_srs_document():
    doc = docx.Document()
    
    # 1. Document Page Setup (1 inch margins)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # 2. Add Cover Page
    for _ in range(3):
        p = doc.add_paragraph()
        format_paragraph(p, space_after=0)
        
    # Title
    p_title = doc.add_paragraph()
    format_paragraph(p_title, space_before=12, space_after=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run_title = p_title.add_run("SOFTWARE REQUIREMENTS SPECIFICATION")
    set_run_font(run_title, name="Calibri Light", size_pt=24, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x79))
    
    # App Name
    p_subtitle = doc.add_paragraph()
    format_paragraph(p_subtitle, space_before=6, space_after=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run_sub = p_subtitle.add_run("Money Manager Java Suite")
    set_run_font(run_sub, name="Calibri", size_pt=18, italic=True, color_rgb=RGBColor(0x2F, 0x55, 0x97))
    
    # Subtitle Details
    p_version = doc.add_paragraph()
    format_paragraph(p_version, space_before=4, space_after=24, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run_ver = p_version.add_run("Version 1.0 • JavaFX Desktop Client & Spring Boot Web Application\nDocument Date: March 26, 2026")
    set_run_font(run_ver, name="Calibri", size_pt=12, color_rgb=RGBColor(89, 89, 89))
    
    # Metadata Box Table (Aligned with actual Java Workspace stack)
    metadata_headers = ["Project Detail", "Specification Value"]
    metadata_data = [
        ["Project Name", "Money Manager Java Suite"],
        ["Document Type", "Software Requirements Specification (SRS)"],
        ["Version", "1.0"],
        ["Language", "Java 21 LTS (Unified Target)"],
        ["Frameworks", "JavaFX 21 (Desktop), Spring Boot 4.0.6 & Thymeleaf (Web)"],
        ["Databases", "SQLite (Desktop local database), PostgreSQL 16 (Web database)"],
        ["Architecture", "3-Tier Architecture + SOLID Principles (JDBC repositories)"],
        ["Status", "Approved — Ready for Production Development"]
    ]
    create_styled_table(doc, metadata_headers, metadata_data, col_widths=[2.5, 4.0])
    
    p_bot = doc.add_paragraph()
    format_paragraph(p_bot, space_before=48)
    
    # Page Break after Cover
    doc.add_page_break()
    
    # 3. Table of Contents Page
    add_styled_heading(doc, "Table of Contents", level=1)
    
    toc_data = [
        ["1.", "Introduction & Purpose", "Page 3"],
        ["2.", "Overall Description", "Page 4"],
        ["3.", "Stakeholders & User Classes", "Page 4"],
        ["4.", "System Architecture", "Page 5"],
        ["5.", "SOLID Design Principles", "Page 6"],
        ["6.", "Functional Requirements", "Page 7"],
        ["7.", "Non-Functional Requirements", "Page 8"],
        ["8.", "Input / Output Specification", "Page 9"],
        ["9.", "Use Case Descriptions", "Page 10"],
        ["10.", "System Diagrams", "Page 11"],
        ["11.", "Database Design", "Page 12"],
        ["12.", "Sequence Diagrams", "Page 13"],
        ["13.", "Future Enhancements", "Page 14"],
        ["14.", "Constraints & Assumptions", "Page 15"],
        ["15.", "Revision History", "Page 15"]
    ]
    
    for item in toc_data:
        p_toc = doc.add_paragraph()
        format_paragraph(p_toc, space_before=2, space_after=3)
        run_num = p_toc.add_run(f"{item[0]}  ")
        set_run_font(run_num, bold=True, color_rgb=RGBColor(0x1F, 0x4E, 0x79))
        run_name = p_toc.add_run(f"{item[1]}")
        set_run_font(run_name)
        dots_count = 65 - len(item[1])
        run_dots = p_toc.add_run(" ." * dots_count)
        set_run_font(run_dots, color_rgb=RGBColor(190, 190, 190))
        run_page = p_toc.add_run(f"  {item[2]}")
        set_run_font(run_page, bold=True)
        
    doc.add_page_break()
    
    # 4. Section 1: Introduction & Purpose
    add_styled_heading(doc, "1. Introduction & Purpose", level=1)
    
    add_styled_heading(doc, "1.1 Purpose", level=2)
    p_purp = doc.add_paragraph()
    format_paragraph(p_purp)
    run_p = p_purp.add_run(
        "This Software Requirements Specification (SRS) document formally describes the functional and non-functional "
        "requirements for the Money Manager Java Suite — a Java-based desktop and web personal finance tool. It serves "
        "as the authoritative contract between stakeholders, developers, QA engineers, and future maintainers. All "
        "design and implementation decisions must conform to this document."
    )
    set_run_font(run_p)
    
    add_styled_heading(doc, "1.2 Scope", level=2)
    p_scope = doc.add_paragraph()
    format_paragraph(p_scope)
    run_s = p_scope.add_run(
        "The Money Manager Suite enables individuals to track income and expenses, categorize transactions, "
        "manage monthly budgets with warnings, set personal financial savings goals, write notes, and view visual report "
        "dashboards. The system is composed of an offline-first JavaFX desktop application utilizing SQLite local storage "
        "and a web application powered by Spring Boot, Thymeleaf templates, and PostgreSQL. It utilizes a strict 3-Tier + "
        "SOLID architecture to ensure absolute maintainability and testability."
    )
    set_run_font(run_s)
    
    add_styled_heading(doc, "1.3 Definitions & Acronyms", level=2)
    def_headers = ["Acronym / Term", "Definition"]
    def_data = [
        ["SRS", "Software Requirements Specification"],
        ["SOLID", "Single Responsibility, Open/Closed, Liskov Substitution, Interface Segregation, Dependency Inversion principles"],
        ["GUI / Client", "Graphical User Interface representing the visual frontends of the application"],
        ["CRUD", "Create, Read, Update, and Delete database operations"],
        ["DTO", "Data Transfer Object (e.g. TransactionDTO, used to exchange structured data between layers safely)"],
        ["JavaFX", "Native Java GUI desktop application framework using FXML and styling skins"],
        ["Spring Boot", "Centralized, lightweight Java framework mapping HTTP routes and managing context beans (Web)"],
        ["Thymeleaf", "Java server-side template engine parsing dynamic XML/HTML5 structures into web browsers"],
        ["SQLite", "Offline-first, self-contained local relational SQL database engine embedded via JDBC (Desktop)"],
        ["PostgreSQL", "Enterprise-scale relational SQL database server providing high-concurrency storage (Web)"],
        ["JDBC", "Java Database Connectivity API used to execute direct SQL and read results safely via PreparedStatements"],
        ["Lombok", "Java library that automatically generates getters, setters, constructors, and loggers during build compilation"]
    ]
    create_styled_table(doc, def_headers, def_data, col_widths=[2.0, 4.5])
    
    # 5. Section 2: Overall Description
    add_styled_heading(doc, "2. Overall Description", level=1)
    
    add_styled_heading(doc, "2.1 Product Perspective", level=2)
    p_pers = doc.add_paragraph()
    format_paragraph(p_pers)
    run_pers = p_pers.add_run(
        "The Money Manager Suite provides two primary deployment interfaces built on shared domain structures and service rules. "
        "The first is an offline-first Desktop Application compiled using Java 21, which writes its records directly to a local "
        "SQLite database (`money-manager.db`). The second is a Web Application also compiled in Java 21, using Spring Boot 4.0.6, "
        "Thymeleaf templates, and connecting to a central PostgreSQL database. This separation guarantees local privacy "
        "for single workstation users while providing central access, responsive view layouts, and concurrent sync for web environments."
    )
    set_run_font(run_pers)
    
    add_styled_heading(doc, "2.2 Product Features", level=2)
    add_bullet_point(doc, "Secure user profiles utilizing passwords safely hashed via BCrypt encryption algorithms.", "Secure Authentication — ")
    add_bullet_point(doc, "Log, edit, and delete transactions categorized in income/expense types.", "Transaction Management — ")
    add_bullet_point(doc, "Set category caps per month, with progress indicators warning at 80% (amber) and alert at 100% (red).", "Smart Budget Limits — ")
    add_bullet_point(doc, "Interactive dashboard charts (PieChart, BarChart) rendering spending totals and balance trends.", "Dashboard Metrics — ")
    add_bullet_point(doc, "Savings milestones with deadlined fund requirements, contribution trackers, and progress states.", "Savings Goals — ")
    add_bullet_point(doc, "Embedded scratchnotes list enabling users to save financial outlines or logs.", "Financial Notes — ")
    add_bullet_point(doc, "Strict validation rules preventing expenses from exceeding defined monthly caps or category limits.", "Validation Checks — ")
    
    add_styled_heading(doc, "2.3 User Needs", level=2)
    p_needs = doc.add_paragraph()
    format_paragraph(p_needs)
    run_n = p_needs.add_run(
        "Users require a modern, secure personal accounting system that functions either completely offline (Desktop Client) "
        "or globally on multiple screens (Web App). Users need simple, fast logging interfaces, graphical trend feedback, and "
        "interactive over-limit budget alerts without setup complications."
    )
    set_run_font(run_n)
    
    # 6. Section 3: Stakeholders & User Classes
    add_styled_heading(doc, "3. Stakeholders & User Classes", level=1)
    p_stake_desc = doc.add_paragraph()
    format_paragraph(p_stake_desc)
    run_s_desc = p_stake_desc.add_run(
        "The primary actors and interested parties in the development and deployment of the Money Manager Java Suite include:"
    )
    set_run_font(run_s_desc)
    
    stake_headers = ["Stakeholder / User Class", "Role in Project", "Primary Interests & Focus Area"]
    stake_data = [
        ["End User", "Primary consumer of the app", "Ease of use, responsive tables, secure offline files (SQLite) or web access (PostgreSQL)."],
        ["Developer", "Implements features and logic", "3-tier separation, constructor dependency injection, JUnit test setups, Lombok builders."],
        ["QA Engineer", "Validates operational correctness", "Precise business rules, validation checking, and automated integration test setups."],
        ["Architect", "Owns technical design & structure", "SOLID compliance, JDBC abstraction patterns, Java 21 LTS platform compatibilities."],
        ["Product Owner", "Defines business requirements & scope", "Feature parity between desktop and web, user retention, timeline, and high UX aesthetic standards."]
    ]
    create_styled_table(doc, stake_headers, stake_data, col_widths=[1.8, 2.2, 2.5])
    
    doc.add_page_break()
    
    # 7. Section 4: System Architecture
    add_styled_heading(doc, "4. System Architecture", level=1)
    p_arch_desc = doc.add_paragraph()
    format_paragraph(p_arch_desc)
    run_a_desc = p_arch_desc.add_run(
        "The entire suite complies with a strict 3-Tier Architecture separating UI structures, business validations, "
        "and SQL data access layers. High-level business services rely purely on repository interfaces, completely isolating "
        "database technology details from controllers and views."
    )
    set_run_font(run_a_desc)
    
    add_styled_heading(doc, "A. Presentation Layer (GUI / Web Views)", level=2)
    add_bullet_point(doc, "Desktop client: JavaFX controllers mapping `.fxml` layout structures, skinned using modular `.css` styles.")
    add_bullet_point(doc, "Web client: Spring WebMVC controllers (`WebAuthController`, `WebDashboardController`) returning dynamic Thymeleaf views.")
    add_bullet_point(doc, "Zero business, calculation, or validation logic resides in this tier — all events delegate directly to the service layer.")
    
    add_styled_heading(doc, "B. Business Logic Layer (Service Layer)", level=2)
    add_bullet_point(doc, "Implements calculations, authentication wrappers, database triggers, and security question checks.")
    add_bullet_point(doc, "Examples: `TransactionService`, `BudgetService`, `AuthService`, `GoalService`, `MonthlyIncomeService`.")
    add_bullet_point(doc, "Completely decoupled from frontends. Exposes parameters purely via lightweight Data Transfer Objects (`TransactionDTO`).")
    add_bullet_point(doc, "Depends only on repository interfaces, allowing dependencies to be easily mocked or replaced.")
    
    add_styled_heading(doc, "C. Data Layer (Repository Layer)", level=2)
    add_bullet_point(doc, "Manages connections via `DatabaseConfig` properties, using plain JDBC with PreparedStatements (no ORM overhead).")
    add_bullet_point(doc, "Defines repository interfaces: `ITransactionRepo`, `IBudgetRepo`, `IGoalRepo`, `INoteRepo`, `IMonthlyBalanceRepo`, `IUserSettingsRepo`.")
    add_bullet_point(doc, "Concrete repository implementations (`JdbcTransactionRepo`, `JdbcUserRepo`) write SQL queries directly.")
    add_bullet_point(doc, "Interchangeable data storage definitions: Desktop connects to local SQLite, and Web connects to PostgreSQL.")
    
    # Architecture Diagram Reference
    add_styled_heading(doc, "Figure 1 — 3-Tier Architecture Overview", level=2)
    create_callout_box(
        doc,
        "┌────────────────────────────────────────────────────────────────────────┐\n"
        "│                          PRESENTATION LAYER                            │\n"
        "│  - Desktop: JavaFX (FXML, CSS, Controllers)                            │\n"
        "│  - Web: Spring Boot Controllers (WebDashboardController, Thymeleaf)     │\n"
        "└───────────────────────────────────┬────────────────────────────────────┘\n"
        "                                    │ Calls (DTOs)\n"
        "┌───────────────────────────────────▼────────────────────────────────────┐\n"
        "│                        BUSINESS LOGIC LAYER                            │\n"
        "│  - AuthService • TransactionService • BudgetService • GoalService     │\n"
        "└───────────────────────────────────┬────────────────────────────────────┘\n"
        "                                    │ Abstractions (Interfaces)\n"
        "┌───────────────────────────────────▼────────────────────────────────────┐\n"
        "│                          REPOSITORY LAYER                              │\n"
        "│  - Interfaces: ITransactionRepo • IBudgetRepo • IUserRepo • IGoalRepo  │\n"
        "│  - Concrete implementations: JdbcTransactionRepo • JdbcBudgetRepo      │\n"
        "└───────────────────┬────────────────────────────────┬───────────────────┘\n"
        "                    │ (SQLite JDBC)                  │ (Postgres Driver)\n"
        "            ┌───────▼───────┐                ┌───────▼───────┐\n"
        "            │ SQLite Local  │                │ Centralized   │\n"
        "            │  Database     │                │ PostgreSQL DB │\n"
        "            └───────────────┘                └───────────────┘"
    )
    
    # Load raw PlantUML code from docs/app.puml
    app_puml_code = read_puml_file("docs/app.puml")
    add_styled_heading(doc, "PlantUML Specification for Architecture (docs/app.puml)", level=3)
    create_code_block(doc, app_puml_code)
    
    doc.add_page_break()
    
    # 8. Section 5: SOLID Design Principles
    add_styled_heading(doc, "5. SOLID Design Principles", level=1)
    p_solid_desc = doc.add_paragraph()
    format_paragraph(p_solid_desc)
    run_so_desc = p_solid_desc.add_run(
        "Both the desktop client and web backend strictly adhere to the SOLID object-oriented patterns "
        "to ensure modularity, high code coverage, and seamless class swappability:"
    )
    set_run_font(run_so_desc)
    
    add_styled_heading(doc, "S — Single Responsibility Principle (SRP)", level=2)
    p_srp = doc.add_paragraph()
    format_paragraph(p_srp)
    run_srp = p_srp.add_run(
        "Each component possesses a single focus of change. For example, AuthService manages credential hashing "
        "and security verification only; JdbcTransactionRepo handles physical SQL transaction querying only; "
        "and AlertHelper is strictly concerned with rendering system alert dialogs."
    )
    set_run_font(run_srp)
    
    add_styled_heading(doc, "O — Open / Closed Principle (OCP)", level=2)
    p_ocp = doc.add_paragraph()
    format_paragraph(p_ocp)
    run_ocp = p_ocp.add_run(
        "Classes are open for structural extension but closed for modification. New repository storage backends "
        "(e.g., swapping local JDBC repositories for standard Spring Data JPA or cloud REST APIs) can be "
        "introduced by declaring new repository classes implementing shared repo interfaces, requiring zero edits inside services."
    )
    set_run_font(run_ocp)
    
    add_styled_heading(doc, "L — Liskov Substitution Principle (LSP)", level=2)
    p_lsp = doc.add_paragraph()
    format_paragraph(p_lsp)
    run_lsp = p_lsp.add_run(
        "Implementations are completely interchangeable with their parent interface boundaries. Any module "
        "interacting with `ITransactionRepo` will execute identically whether it is injected with `JdbcTransactionRepo` "
        "referencing SQLite or a mock repository class during test runs."
    )
    set_run_font(run_lsp)
    
    add_styled_heading(doc, "I — Interface Segregation Principle (ISP)", level=2)
    p_isp = doc.add_paragraph()
    format_paragraph(p_isp)
    run_isp = p_isp.add_run(
        "Interfaces are kept clean, modular, and highly atomic. Relational tables map to distinct interfaces: "
        "IBudgetRepo, ITransactionRepo, IGoalRepo, INoteRepo, and IUserSettingsRepo. Classes are never forced "
        "to implement service hookups or persistence query methods that they do not explicitly require."
    )
    set_run_font(run_isp)
    
    add_styled_heading(doc, "D — Dependency Inversion Principle (DIP)", level=2)
    p_dip = doc.add_paragraph()
    format_paragraph(p_dip)
    run_dip = p_dip.add_run(
        "High-level service classes depend purely on repository interface definitions, never on concrete SQL classes. "
        "All connections and sub-repositories are injected dynamically at runtime via class constructors (Constructor Injection), "
        "facilitating straightforward unit test mocking using JUnit 5."
    )
    set_run_font(run_dip)
    
    doc.add_page_break()
    
    # 9. Section 6: Functional Requirements
    add_styled_heading(doc, "6. Functional Requirements", level=1)
    
    add_styled_heading(doc, "FR-01: Transaction Management", level=2)
    fr1_headers = ["Requirement ID", "Functional Requirement Description"]
    fr1_data = [
        ["FR-01.1", "The user shall be able to record new transactions specifying: title, amount, category, type (income/expense), and custom date."],
        ["FR-01.2", "The user shall be able to view, query, and filter transactions dynamically based on date range, specific category, and type."],
        ["FR-01.3", "The user shall be able to modify any parameter (name, date, category, amount) of existing transaction records safely."],
        ["FR-01.4", "The user shall be able to delete any logged transaction, accompanied by a double-confirmation visual alert prompt."],
        ["FR-01.5", "The system shall allow transaction entry parameters to feed directly into monthly budget calculation updates in real time."]
    ]
    create_styled_table(doc, fr1_headers, fr1_data, col_widths=[1.5, 5.0])
    
    add_styled_heading(doc, "FR-02: Budget Management", level=2)
    fr2_headers = ["Requirement ID", "Functional Requirement Description"]
    fr2_data = [
        ["FR-02.1", "The user shall define specific monthly budgets on a per-category basis with defined maximum spending limits."],
        ["FR-02.2", "The system shall dynamically aggregate expenditures and compare totals against active budget limits in real time."],
        ["FR-02.3", "The presentation layer shall display color-coded warning limits: progress bars will render green (<80%), shift to amber/orange (80-99%) to warn, and shift to red (>=100%) to indicate budget breaches."],
        ["FR-02.4", "The user shall be able to create, edit, and clear monthly category budgets easily."]
    ]
    create_styled_table(doc, fr2_headers, fr2_data, col_widths=[1.5, 5.0])
    
    add_styled_heading(doc, "FR-03: Dashboard & Analytics Reports", level=2)
    fr3_headers = ["Requirement ID", "Functional Requirement Description"]
    fr3_data = [
        ["FR-03.1", "The UI shall calculate and display KPI summary blocks representing Total Income, Total Expenses, and Net Balance for the current calendar month."],
        ["FR-03.2", "The UI shall render a graphical PieChart showing categorical expense percentages relative to total outgoing spend."],
        ["FR-03.3", "The UI shall render a graphical BarChart illustrating expense trends and month-on-month balance changes over a 6-month window."],
        ["FR-03.4", "The system shall integrate services facilitating data exports and generating reports for users."]
    ]
    create_styled_table(doc, fr3_headers, fr3_data, col_widths=[1.5, 5.0])
    
    add_styled_heading(doc, "FR-04: Notes & Savings Goals", level=2)
    fr4_headers = ["Requirement ID", "Functional Requirement Description"]
    fr4_data = [
        ["FR-04.1", "The user shall be able to create and manage personal notes specifying a title, text content, and timestamp."],
        ["FR-04.2", "The user shall define long-term savings goals with title, target amount, current saved progress, and target deadline dates."],
        ["FR-04.3", "The system shall calculate percentage progress and display dynamic bars representing goal completion milestones."],
        ["FR-04.4", "The system shall evaluate deadlines and flag targets color-coded: green (on track), amber (deadline within 30 days and progress <80%), and red (past deadline and incomplete)."]
    ]
    create_styled_table(doc, fr4_headers, fr4_data, col_widths=[1.5, 5.0])
    
    doc.add_page_break()
    
    # 10. Section 7: Non-Functional Requirements
    add_styled_heading(doc, "7. Non-Functional Requirements", level=1)
    
    nfr_headers = ["ID", "Category", "Non-Functional Requirement Specification"]
    nfr_data = [
        ["NFR-01", "Performance", "The desktop client shall boot to the login screen in under 3 seconds. The web interface shall load views and map routes within 200ms of user input."],
        ["NFR-02", "Reliability", "Database transactions must execute atomically. Core services must handle SQL exceptions and rollback partially written records to prevent corruption."],
        ["NFR-03", "Usability", "Interfaces (both JavaFX and Thymeleaf) must be highly intuitive, letting new users perform transaction logs and see budget changes within 5 minutes without documentation."],
        ["NFR-04", "Security", "All user passwords must be securely hashed and verified utilizing BCrypt algorithms (`at.favre.lib:bcrypt` package). Plaintext passwords must never be stored."],
        ["NFR-05", "Maintainability", "Unit tests (using JUnit 5) must exceed 80% coverage across core services. All public classes and service methods must feature descriptive Javadocs."],
        ["NFR-06", "Portability", "Desktop client must build into platform-specific fat JARs running on any system with Java 21 installed. Web application must deploy on any Java 21 compatible runtime."],
        ["NFR-07", "Scalability", "Repositories must be fully decoupled via interfaces, allowing the system to swap SQLite JDBC connectors for PostgreSQL or Cloud APIs without changing services."],
        ["NFR-08", "Data Integrity", "All database tables must strictly enforce foreign key rules and cascading delete parameters (`ON DELETE CASCADE`) to prevent orphaned transaction rows."]
    ]
    create_styled_table(doc, nfr_headers, nfr_data, col_widths=[0.8, 1.2, 4.5])
    
    doc.add_page_break()
    
    # 11. Section 8: Input / Output Specification
    add_styled_heading(doc, "8. Input / Output Specification", level=1)
    
    p_io_desc = doc.add_paragraph()
    format_paragraph(p_io_desc)
    run_io_desc = p_io_desc.add_run(
        "The following specifications map how input metrics are consumed by services and converted "
        "into graphical dashboards, files, and progress notifications:"
    )
    set_run_font(run_io_desc)
    
    add_styled_heading(doc, "Figure 2 — Input / Output Flow Diagram", level=2)
    create_callout_box(
        doc,
        "INPUT PARAMETERS                             JAVA CORE SERVICE LAYER                     OUTPUT INTERFACES\n"
        "┌──────────────────────┐                 ┌─────────────────────────────┐                 ┌──────────────────────┐\n"
        "│  Transaction Entry   │ ──────────────> │                             │ ──────────────> │ Dashboard KPIs       │\n"
        "│  (Amount, Date, Cat) │                 │                             │                 │ (Total Income, Net)  │\n"
        "├──────────────────────┤                 │  AuthService                │                 ├──────────────────────┤\n"
        "│  Category Budget     │ ──────────────> │  TransactionService         │ ──────────────> │ Visual Alerts        │\n"
        "│  (Amount Cap, Month) │                 │  BudgetService              │                 │ (Amber/Red Indicators)│\n"
        "├──────────────────────┤                 │  GoalService                │                 ├──────────────────────┤\n"
        "│  Savings Goals       │ ──────────────> │  MonthlyIncomeService       │ ──────────────> │ Rendered Views       │\n"
        "│  (Target, Deadline)  │                 │                             │                 │ (JavaFX / Thymeleaf) │\n"
        "├──────────────────────┤                 │                             │                 ├──────────────────────┤\n"
        "│  Notes & Settings    │ ──────────────> │                             │ ──────────────> │ Goal Completion      │\n"
        "│  (Monthly Income)    │                 └─────────────────────────────┘                 │ (% Progress Bars)    │\n"
        "└──────────────────────┘                                                                 └──────────────────────┘"
    )
    
    add_styled_heading(doc, "Key Data Inputs", level=2)
    in_headers = ["Input Source", "Format", "Primary Validation Rules"]
    in_data = [
        ["Transaction Details", "User Form / Request Parameters", "Amount must be a positive decimal (> 0). Transaction date must be specified and not in the future."],
        ["Budget Threshold", "User Form / Request Parameters", "Limit cap must be positive (> 0). Category descriptor must match a valid database reference."],
        ["Savings Goal", "User Form / Request Parameters", "Goal name is mandatory. Target amount must be positive, and deadline date must be in the future."],
        ["Filter Parameters", "GUI Pickers / Query Parameters", "Start date must be less than or equal to the end date range."],
        ["User Preferences", "Settings Panel / SQL fields", "Monthly income parameter must be zero or positive. Settings tied to user account ID."]
    ]
    create_styled_table(doc, in_headers, in_data, col_widths=[2.0, 2.0, 2.5])
    
    # 12. Section 9: Use Case Descriptions
    add_styled_heading(doc, "9. Use Case Descriptions", level=1)
    
    add_styled_heading(doc, "Figure 3 — System Use Case Diagram", level=2)
    usecase_puml = read_puml_file("docs/use-case.puml")
    create_code_block(doc, usecase_puml)
    
    add_styled_heading(doc, "UC-01: Add Transaction", level=2)
    p_uc1_act = doc.add_paragraph()
    format_paragraph(p_uc1_act, space_after=2)
    run_act = p_uc1_act.add_run("Primary Actor: ")
    set_run_font(run_act, bold=True)
    p_uc1_act.add_run("End User")
    
    p_uc1_pre = doc.add_paragraph()
    format_paragraph(p_uc1_pre, space_after=2)
    run_pre = p_uc1_pre.add_run("Preconditions: ")
    set_run_font(run_pre, bold=True)
    p_uc1_pre.add_run("User is logged in; active session (Web) or active state (Desktop) is valid.")
    
    p_uc1_flow = doc.add_paragraph()
    format_paragraph(p_uc1_flow, space_after=2)
    run_flow = p_uc1_flow.add_run("Main Flow of Events:")
    add_bullet_point(doc, "User navigates to the Transactions view.")
    add_bullet_point(doc, "User clicks 'Add' or enters details inside the transaction form.")
    add_bullet_point(doc, "System displays or maps fields for Title, Amount, Category, Date, and Type.")
    add_bullet_point(doc, "User submits the transaction details.")
    add_bullet_point(doc, "System triggers service-tier input validation checks.")
    add_bullet_point(doc, "If the transaction is an 'EXPENSE', `BudgetService` checks category caps and monthly balances.")
    add_bullet_point(doc, "If validations pass, `TransactionService` invokes `ITransactionRepo.save()` to write the record.")
    add_bullet_point(doc, "System refreshes the table grid UI and updates dashboard graphs.")
    
    p_uc1_alt = doc.add_paragraph()
    format_paragraph(p_uc1_alt, space_after=2)
    run_alt = p_uc1_alt.add_run("Alternate Flow:")
    set_run_font(run_alt, bold=True)
    add_bullet_point(doc, "Step 5a: Validation checks fail (e.g., negative amount, future date). The system halts the operation, displays an alert dialog (Desktop) or dynamic error message (Web), and prompts the user to correct the fields.")
    add_bullet_point(doc, "Step 6a: Budget limit checks indicate that the new expense will exceed category caps. The service layer triggers an exception, alerting the user about the budget breach.")
    
    p_uc1_post = doc.add_paragraph()
    format_paragraph(p_uc1_post)
    run_post = p_uc1_post.add_run("Postconditions: ")
    set_run_font(run_post, bold=True)
    p_uc1_post.add_run("The transaction is committed to the database; monthly balances and dashboard indicators recalculate immediately.")
    
    # UC-02
    add_styled_heading(doc, "UC-02: Set & Monitor Budget", level=2)
    p_uc2_act = doc.add_paragraph()
    format_paragraph(p_uc2_act, space_after=2)
    run_act2 = p_uc2_act.add_run("Primary Actor: ")
    set_run_font(run_act2, bold=True)
    p_uc2_act.add_run("End User")
    
    p_uc2_pre = doc.add_paragraph()
    format_paragraph(p_uc2_pre, space_after=2)
    run_pre2 = p_uc2_pre.add_run("Preconditions: ")
    set_run_font(run_pre2, bold=True)
    p_uc2_pre.add_run("User is authenticated; target category exists in the system defaults.")
    
    p_uc2_flow = doc.add_paragraph()
    format_paragraph(p_uc2_flow, space_after=2)
    run_flow2 = p_uc2_flow.add_run("Main Flow of Events:")
    add_bullet_point(doc, "User navigates to the Budgets view.")
    add_bullet_point(doc, "User selects a category and defines a monthly spending cap limit.")
    add_bullet_point(doc, "System validates parameters and inserts or updates the budget record.")
    add_bullet_point(doc, "When transaction expenses are added, the system matches their category and date to active budgets.")
    add_bullet_point(doc, "System displays progress percentages graphically per budget row.")
    
    p_uc2_alt = doc.add_paragraph()
    format_paragraph(p_uc2_alt, space_after=2)
    run_alt2 = p_uc2_alt.add_run("Alternate Flow:")
    set_run_font(run_alt2, bold=True)
    add_bullet_point(doc, "Step 5a: Outgoing spend reaches 80% to 99% of the budget cap limit. The UI renders the progress bar in amber/orange.")
    add_bullet_point(doc, "Step 5b: Outgoing spend reaches or exceeds 100% of the budget cap limit. The UI renders the progress bar in red, and flashes a budget over-limit alert warning.")
    
    p_uc2_post = doc.add_paragraph()
    format_paragraph(p_uc2_post)
    run_post2 = p_uc2_post.add_run("Postconditions: ")
    set_run_font(run_post2, bold=True)
    p_uc2_post.add_run("The category budget constraints are configured; color-coded safety margins monitor active transactions.")
    
    doc.add_page_break()
    
    # 13. Section 10: System Diagrams
    add_styled_heading(doc, "10. System Diagrams", level=1)
    
    add_styled_heading(doc, "10.1 Entity-Relationship Diagram", level=2)
    p_erd_d = doc.add_paragraph()
    format_paragraph(p_erd_d)
    run_erd_d = p_erd_d.add_run(
        "The logical entity-relationship diagram below maps all database tables. To ensure data isolation, "
        "every transaction, budget, savings goal, and setting record holds a foreign key reference linked directly "
        "to a parent User account."
    )
    set_run_font(run_erd_d)
    
    add_styled_heading(doc, "Figure 4 — Entity-Relationship Diagram", level=2)
    erd_puml = read_puml_file("docs/erd.puml")
    create_code_block(doc, erd_puml)
    
    # 14. Section 11: Database Design (Comparing both SQLite & PostgreSQL physical schemas)
    add_styled_heading(doc, "11. Database Design", level=1)
    
    p_db_d = doc.add_paragraph()
    format_paragraph(p_db_d)
    run_db_d = p_db_d.add_run(
        "The Money Manager Suite supports two physical databases: SQLite for local offline desktop runs, and "
        "PostgreSQL for centralized web application deployments. The tables share identical logical definitions, "
        "but utilize platform-specific constraints and data types mapped below:"
    )
    set_run_font(run_db_d)
    
    # Table users
    add_styled_heading(doc, "Table: users", level=2)
    u_headers = ["Column Name", "SQLite Type (Desktop)", "PostgreSQL Type (Web)", "Key Constraints / Details"]
    u_data = [
        ["user_id", "INTEGER", "BIGSERIAL", "PRIMARY KEY, AUTOINCREMENT, Unique identifier"],
        ["username", "TEXT", "VARCHAR(50)", "NOT NULL, UNIQUE, Account lookup name"],
        ["password_hash", "TEXT", "TEXT", "NOT NULL, BCrypt hashed password credentials"],
        ["security_question", "TEXT", "N/A (Desktop Only)", "Nullable, desktop profile recovery query"],
        ["security_answer_hash", "TEXT", "N/A (Desktop Only)", "Nullable, desktop hashed answer profile"],
        ["created_at", "TEXT", "TIMESTAMPTZ", "NOT NULL, Default now() or ISO 8601 string"]
    ]
    create_styled_table(doc, u_headers, u_data, col_widths=[1.5, 1.5, 1.5, 2.0])
    
    # Table transactions
    add_styled_heading(doc, "Table: transactions", level=2)
    tx_headers = ["Column Name", "SQLite Type (Desktop)", "PostgreSQL Type (Web)", "Key Constraints / Details"]
    tx_data = [
        ["transaction_id", "INTEGER", "BIGSERIAL", "PRIMARY KEY, AUTOINCREMENT"],
        ["user_id", "INTEGER", "BIGINT", "FOREIGN KEY references users(user_id) ON DELETE CASCADE"],
        ["name", "TEXT", "VARCHAR(100)", "NOT NULL, transaction title"],
        ["amount", "REAL", "NUMERIC(12,2)", "NOT NULL, check constraint (amount > 0)"],
        ["category", "TEXT", "VARCHAR(50)", "NOT NULL, category descriptor"],
        ["tx_type", "TEXT", "VARCHAR(10)", "NOT NULL, check constraint in ('INCOME', 'EXPENSE')"],
        ["tx_date", "TEXT", "DATE", "NOT NULL, transaction execution date"],
        ["created_at", "TEXT", "TIMESTAMPTZ", "NOT NULL, creation system timestamp"]
    ]
    create_styled_table(doc, tx_headers, tx_data, col_widths=[1.5, 1.5, 1.5, 2.0])
    
    # Table budgets
    add_styled_heading(doc, "Table: budgets", level=2)
    b_headers = ["Column Name", "SQLite Type (Desktop)", "PostgreSQL Type (Web)", "Key Constraints / Details"]
    b_data = [
        ["budget_id", "INTEGER", "BIGSERIAL", "PRIMARY KEY, AUTOINCREMENT"],
        ["user_id", "INTEGER", "BIGINT", "FOREIGN KEY references users(user_id) ON DELETE CASCADE"],
        ["category", "TEXT", "VARCHAR(50)", "NOT NULL, category description"],
        ["amount_cap", "REAL", "NUMERIC(12,2)", "NOT NULL, check constraint (amount_cap > 0)"],
        ["month", "INTEGER", "SMALLINT", "NOT NULL, check constraint (month between 1 and 12)"],
        ["year", "INTEGER", "SMALLINT", "NOT NULL, check constraint (year >= 2020)"],
        ["CONSTRAINT", "UNIQUE(composite)", "UNIQUE(composite)", "Composite UNIQUE constraints on (user_id, category, month, year)"]
    ]
    create_styled_table(doc, b_headers, b_data, col_widths=[1.5, 1.5, 1.5, 2.0])
    
    # Table savings_goals
    add_styled_heading(doc, "Table: Goals (savings_goals)", level=2)
    g_headers = ["Column Name", "SQLite Type (Desktop)", "PostgreSQL Type (Web)", "Key Constraints / Details"]
    g_data = [
        ["goal_id", "INTEGER", "BIGSERIAL", "PRIMARY KEY, AUTOINCREMENT"],
        ["user_id", "INTEGER", "BIGINT", "FOREIGN KEY references users(user_id) ON DELETE CASCADE"],
        ["name", "TEXT", "VARCHAR(100)", "NOT NULL, savings milestone name"],
        ["target_amount", "REAL", "NUMERIC(12,2)", "NOT NULL, check constraint (target_amount > 0)"],
        ["saved_amount", "REAL", "NUMERIC(12,2)", "NOT NULL, DEFAULT 0, check constraint (saved_amount >= 0)"],
        ["deadline", "TEXT", "DATE", "Nullable, targeted complete date limit"],
        ["created_at", "TEXT", "TIMESTAMPTZ", "NOT NULL, Default creation timestamp"]
    ]
    create_styled_table(doc, g_headers, g_data, col_widths=[1.5, 1.5, 1.5, 2.0])
    
    doc.add_page_break()
    
    # 15. Section 12: Sequence Diagrams
    add_styled_heading(doc, "12. Sequence Diagrams", level=1)
    
    add_styled_heading(doc, "12.1 Add Expense Transaction Flow", level=2)
    p_seq_d = doc.add_paragraph()
    format_paragraph(p_seq_d)
    run_seq_d = p_seq_d.add_run(
        "The following sequence diagram outlines execution pathways when the user logs an expense. "
        "It highlights control flowing from the UI/Web Controller layer down through service validations "
        "(limit validation on active budgets and monthly totals) and committing records inside database repository classes."
    )
    set_run_font(run_seq_d)
    
    add_styled_heading(doc, "Figure 5 — Sequence Diagram: Add Transaction (with budget checks)", level=2)
    seq_puml = read_puml_file("docs/sequence-add-expense-transaction.puml")
    create_code_block(doc, seq_puml)
    
    # 16. Section 13: Future Enhancements
    add_styled_heading(doc, "13. Future Enhancements", level=1)
    
    add_styled_heading(doc, "Mobile Client Ports (Android / iOS)", level=2)
    p_f1 = doc.add_paragraph()
    format_paragraph(p_f1)
    run_f1 = p_f1.add_run(
        "Bundle the shared core services and entities into modular, reusable packages. Design a fully customized "
        "mobile client using Flutter or React Native, utilizing REST APIs to achieve real-time synchronization "
        "with the core PostgreSQL web backend."
    )
    set_run_font(run_f1)
    
    add_styled_heading(doc, "Cloud Synchronizer Utilities", level=2)
    p_f2 = doc.add_paragraph()
    format_paragraph(p_f2)
    run_f2 = p_f2.add_run(
        "Provide local desktop users with optional sup-sync backup capabilities. Build a connection client "
        "that securely pushes records from the local offline SQLite database file to a centralized cloud PostgreSQL "
        "server or REST API, using Supabase or AWS RDS hosting instances."
    )
    set_run_font(run_f2)
    
    add_styled_heading(doc, "AI-Driven Spending Insights & Coaches", level=2)
    p_f3 = doc.add_paragraph()
    format_paragraph(p_f3)
    run_f3 = p_f3.add_run(
        "Integrate lightweight local machine learning models (e.g., via ONNX or customized pipelines) to "
        "automatically audit monthly transaction entries, categorize outgoing expenditures based on transaction names, "
        "detect spending anomalies, and generate financial tips."
    )
    set_run_font(run_f3)
    
    add_styled_heading(doc, "Polished Document Export Engines", level=2)
    p_f4 = doc.add_paragraph()
    format_paragraph(p_f4)
    run_f4 = p_f4.add_run(
        "Deploy robust reporting libraries (e.g., OpenPDF or JasperReports for Java) to let users export "
        "highly customized financial summaries directly to styled PDF statements or pre-configured Excel workbooks containing "
        "mathematical formulas."
    )
    set_run_font(run_f4)
    
    add_styled_heading(doc, "Multi-User Profile sharing", level=2)
    p_f5 = doc.add_paragraph()
    format_paragraph(p_f5)
    run_f5 = p_f5.add_run(
        "Expand web authentication rules to support family access controls. Let household profiles register view-only "
        "or collaborative editor permissions to collectively log transactions and see budget limits on unified accounts."
    )
    set_run_font(run_f5)
    
    doc.add_page_break()
    
    # 17. Section 14: Constraints & Assumptions
    add_styled_heading(doc, "14. Constraints & Assumptions", level=1)
    
    add_styled_heading(doc, "Development & Structural Constraints", level=2)
    add_bullet_point(doc, "The initial release v1.0 locks desktop runs to isolated local user profiles; central collaborative profile sync is deferred to v2.0.")
    add_bullet_point(doc, "Desktop GUI client layouts must use pure JavaFX controls and custom CSS styles — no embedded HTML/JS web frameworks inside v1.0 desktop.")
    add_bullet_point(doc, "All desktop data remains completely local inside SQLite files; cloud synchronizations are purely user-initiated and optional.")
    add_bullet_point(doc, "The desktop client must run completely offline without demanding an active internet or API connection.")
    add_bullet_point(doc, "Desktop application builds must package into lightweight, standalone shaded executable JARs using the Maven Shade plugin.")
    
    add_styled_heading(doc, "Core Architectural Assumptions", level=2)
    add_bullet_point(doc, "Users possess standard financial tracking familiarity (e.g., logging income vs expenses, category caps, net balances).")
    add_bullet_point(doc, "The development workstation includes Maven 3.9+ and JDK 21 setups for both desktop and web modules.")
    add_bullet_point(doc, "The SQLite JDBC driver loaded via Maven is sufficient to support standard transaction concurrency on single-user runs.")
    add_bullet_point(doc, "All monetary fields are handled and computed utilizing the BigDecimal class to prevent floating-point accuracy errors. Databases save numerical values as NUMERIC or REAL.")
    add_bullet_point(doc, "Transaction timestamps are saved as ISO 8601 YYYY-MM-DD text fields (SQLite) or native DATE fields (PostgreSQL) to ensure platform-independent date math.")
    
    # 18. Section 15: Revision History
    add_styled_heading(doc, "15. Revision History", level=1)
    
    rev_headers = ["Version", "Revision Date", "Author / Role", "Summary of Document Changes"]
    rev_data = [
        ["1.0", "2026-03-26", "Project Architect", "Initial SRS release — JavaFX desktop client & Spring Boot web application requirements fully completed and verified."],
        ["0.2", "2024-12-01", "Project Architect", "Added detailed SOLID principle definitions and physical database design comparisons for SQLite and PostgreSQL."],
        ["0.1", "2024-11-15", "Product Owner", "Initial draft document — core objectives, stakeholders, and scope definitions established."]
    ]
    create_styled_table(doc, rev_headers, rev_data, col_widths=[0.8, 1.2, 1.8, 2.7])
    
    setup_header_footer(doc)
    
    # Save the polished document
    output_path = "docs/Software_Requirements_Specification.docx"
    try:
        # Create the docs directory if it doesn't exist
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        print(f"[SUCCESS] Java-Centric SRS Document generated and saved to: {output_path}")
    except PermissionError:
        print(f"\n[WARNING] Permission denied when saving to '{output_path}'.")
        print("This usually means the document is currently open in Microsoft Word.")
        print("Please CLOSE the Word document and run this script again to apply your updates!")

if __name__ == "__main__":
    generate_srs_document()
